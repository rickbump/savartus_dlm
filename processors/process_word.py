import os
from docx import Document


def process_word(file_path):
    metadata = {}
    extracted_content = {}
    try:
        # Basic file metadata
        file_metadata = os.stat(file_path)
        metadata['file_size'] = file_metadata.st_size
        metadata['file_modified'] = file_metadata.st_mtime

        # Load Word document
        doc = Document(file_path)

        # Document properties (core properties)
        core_properties = doc.core_properties
        metadata['author'] = core_properties.author
        metadata['title'] = core_properties.title
        metadata['created'] = core_properties.created
        metadata['modified'] = core_properties.modified
        metadata['subject'] = core_properties.subject
        metadata['keywords'] = core_properties.keywords

        # Document structure metadata
        metadata['paragraph_count'] = len(doc.paragraphs)  # Only count paragraphs, no text metadata
        metadata['table_count'] = len(doc.tables)

        # Extract headers and footers
        headers = []
        footers = []

        # Check for headers and footers in different sections
        for section in doc.sections:
            header = section.header
            footer = section.footer

            # Add header/footer text to their respective lists
            headers.append('\n'.join([p.text for p in header.paragraphs]))
            footers.append('\n'.join([p.text for p in footer.paragraphs]))

        metadata['page_headers'] = list(dict.fromkeys(headers))
        metadata['page_footers'] = list(dict.fromkeys(footers))


        # Process all tables
        tables_metadata = []
        for table_index, table in enumerate(doc.tables):
            table_info = {
                'table_index': table_index + 1,
                'row_count': len(table.rows),
                'column_count': len(table.columns),
            }

            # Process merged cells in the table
            merged_cells = []
            for row in table.rows:
                for cell in row.cells:
                    if len(cell._element.xpath('.//w:vMerge')) > 0:
                        merged_cells.append(cell.text)
            table_info['merged_cells'] = merged_cells

            # Store table metadata
            tables_metadata.append(table_info)
        metadata['tables'] = tables_metadata

        # Extract main document text (excluding headers and footers)
        extracted_content['body_text'] = '\n'.join([p.text for p in doc.paragraphs])

        # Optionally: You can now remove header/footer text from the body if it appears within the body text
        for header_text in headers:
            extracted_content['body_text'] = extracted_content['body_text'].replace(header_text, "")

        for footer_text in footers:
            extracted_content['body_text'] = extracted_content['body_text'].replace(footer_text, "")

    except Exception as e:
        metadata['word_error'] = str(e)

    return metadata, extracted_content